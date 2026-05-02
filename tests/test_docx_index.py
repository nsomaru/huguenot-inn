from pathlib import Path
from zipfile import ZipFile

import fitz
from docx import Document

from huguenot.documents import create_authorities_index_docx, create_matter_authorities_index_docx
from huguenot.domain import Court, DocumentHeaderInput, Matter, Party, PartySide, PDFItem, ProceedingType


def make_pdf(path: Path, text: str, pages: int = 1) -> None:
    doc = fitz.open()
    try:
        for _ in range(pages):
            page = doc.new_page()
            page.insert_text((72, 72), text)
        doc.save(path)
    finally:
        doc.close()


def test_plain_authorities_index_uses_page_ranges(tmp_path: Path) -> None:
    pdf = tmp_path / "authority.pdf"
    make_pdf(pdf, "Authority", pages=2)
    output = tmp_path / "index.docx"

    create_authorities_index_docx([PDFItem(pdf, "Authority title")], output)

    doc = Document(str(output))
    assert doc.tables[0].cell(1, 1).text == "Authority title"
    assert doc.tables[0].cell(1, 2).text == "1-2"


def test_matter_authorities_index_contains_header_parties_and_document_title(tmp_path: Path) -> None:
    pdf = tmp_path / "authority.pdf"
    make_pdf(pdf, "Authority", pages=1)
    output = tmp_path / "matter_index.docx"
    matter = Matter(
        court=Court("IN THE HIGH COURT OF SOUTH AFRICA", "(GAUTENG DIVISION, JOHANNESBURG)"),
        proceeding_type=ProceedingType.APPLICATION,
        case_number="2026-086328",
        parties=(
            Party("Axim (Pty) Ltd", PartySide.BRINGING, 1),
            Party("Moodie", PartySide.OPPOSING, 1),
        ),
    )

    create_matter_authorities_index_docx(
        matter,
        DocumentHeaderInput("Respondents' Authorities Bundle"),
        [PDFItem(pdf, "Authority title")],
        output,
    )

    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    assert "IN THE HIGH COURT OF SOUTH AFRICA" in text
    assert "2026-086328" in text
    assert "RESPONDENTS' AUTHORITIES BUNDLE" in text


def test_matter_authorities_index_uses_table_header_rules_bold_case_and_superscript_ordinals(tmp_path: Path) -> None:
    pdf = tmp_path / "authority.pdf"
    make_pdf(pdf, "Authority", pages=2)
    output = tmp_path / "matter_index.docx"
    matter = Matter(
        court=Court("IN THE HIGH COURT OF SOUTH AFRICA", "(GAUTENG DIVISION, JOHANNESBURG)"),
        proceeding_type=ProceedingType.APPLICATION,
        case_number="2026-086328",
        parties=(
            Party("First Applicant", PartySide.BRINGING, 1),
            Party("Second Applicant", PartySide.BRINGING, 2),
            Party("First Respondent", PartySide.OPPOSING, 1),
            Party("Second Respondent", PartySide.OPPOSING, 2),
        ),
    )

    create_matter_authorities_index_docx(
        matter,
        DocumentHeaderInput("Respondents' Authorities Bundle"),
        [PDFItem(pdf, "Authority title")],
        output,
    )

    doc = Document(str(output))
    assert len(doc.tables) >= 2
    parties_table = doc.tables[0]
    assert "FIRST APPLICANT" in parties_table.cell(0, 0).text
    assert "1st Applicant" in parties_table.cell(0, 1).text
    tramline_width = doc.tables[1].columns[0].width
    party_name_width = parties_table.columns[0].width
    party_label_width = parties_table.columns[1].width
    assert tramline_width is not None
    assert party_name_width is not None
    assert party_label_width is not None
    assert tramline_width == party_name_width + party_label_width
    assert doc.tables[-1].cell(1, 2).text == "1-2"

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode()
    assert "\\t" not in xml
    assert "CASE NO:" in xml
    assert "<w:b/>" in xml or "<w:b " in xml
    assert 'w:vertAlign w:val="superscript"' in xml
    assert '<w:tblW w:type="dxa" w:w="10368"/>' in xml
    assert '<w:bottom w:val="single"' in xml
    assert "w:left=" in xml and "w:hanging=" in xml
    assert "w:tcMar" in xml
