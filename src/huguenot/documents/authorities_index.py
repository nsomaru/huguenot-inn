from __future__ import annotations

from collections.abc import Sequence
from pathlib import Path
from typing import Any, cast

import fitz
from docx import Document
from docx.document import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from huguenot.domain import (
    BundleIndexEntry,
    BundleIndexRow,
    BundleListItem,
    DocumentHeaderInput,
    IndexSeparatorEntry,
    Matter,
    PartySide,
    PDFItem,
    build_bundle_index_entries,
    build_bundle_index_rows,
    normalize_flag_colour,
    party_label,
)
from huguenot.domain.legal_titles import normalize_legal_display_title

DEFAULT_INDEX_FONT = "Times New Roman"
PARTIES_TABLE_WIDTH_INCHES = 7.2
PARTIES_TABLE_COLUMN_WIDTHS = (5.4, 1.8)
TABLE_CELL_MARGIN_TWIPS = 100
HEADING_TRAMLINE_SPACING_PT = 9


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, bold: bool = False, align=None, *, font_name: str = DEFAULT_INDEX_FONT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def get_pdf_page_count(path: Path) -> int:
    source = fitz.open(path)
    try:
        return source.page_count
    finally:
        source.close()


def get_index_entries(
    pdf_items: Sequence[PDFItem],
    *,
    start_page: int = 1,
    flag_colours: Sequence[str] | None = None,
) -> list[BundleIndexEntry]:
    return build_bundle_index_entries(
        pdf_items,
        get_page_count=lambda item: get_pdf_page_count(item.path),
        start_page=start_page,
        flag_colours=flag_colours,
    )


def get_index_rows(
    bundle_items: Sequence[BundleListItem],
    *,
    start_page: int = 1,
    flag_colours: Sequence[str] | None = None,
) -> list[BundleIndexRow]:
    return build_bundle_index_rows(
        bundle_items,
        get_page_count=lambda item: get_pdf_page_count(item.path),
        start_page=start_page,
        flag_colours=flag_colours,
    )


def configure_document(doc: WordDocument, *, font_name: str = DEFAULT_INDEX_FONT) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = cast(Any, doc.styles["Normal"])
    style.font.name = font_name
    style.font.size = Pt(12)


def set_hanging_indent(paragraph, *, indent_twips: int = 180, hanging_twips: int = 180) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(indent_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))


def set_cell_margins(
    cell,
    *,
    top: int = TABLE_CELL_MARGIN_TWIPS,
    start: int = TABLE_CELL_MARGIN_TWIPS,
    bottom: int = TABLE_CELL_MARGIN_TWIPS,
    end: int = TABLE_CELL_MARGIN_TWIPS,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths: tuple[float, ...]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            cell.width = Inches(width)
            set_cell_margins(cell)


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_borders(cell, *, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if edge == "bottom" and bottom:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), "000000")
        else:
            node.set(qn("w:val"), "nil")


def set_cell_right_border(cell, colour_hex: str, *, size: int = 36) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    colour = normalize_flag_colour(colour_hex).lstrip("#")
    for edge in ("end", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), colour)


def add_authorities_table(
    doc: WordDocument,
    entries: Sequence[BundleIndexRow],
    *,
    font_name: str = DEFAULT_INDEX_FONT,
    colour_page_ranges: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(5.7)
    table.columns[2].width = Inches(1.25)

    header_cells = table.rows[0].cells
    set_cell_text(header_cells[0], "No", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
    set_cell_text(header_cells[1], "Item", bold=True, font_name=font_name)
    set_cell_text(header_cells[2], "Page no", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
    set_repeat_table_header(table.rows[0])

    for entry in entries:
        if isinstance(entry, IndexSeparatorEntry):
            row = table.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(18)
            row_cells = row.cells
            set_cell_text(row_cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
            set_cell_text(
                row_cells[1],
                entry.title,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                font_name=font_name,
            )
            set_cell_text(row_cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
            continue
        number = entry.item_number
        item = entry.item
        page_range = entry.page_range
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(18)
        row_cells = row.cells
        set_cell_text(row_cells[0], str(number), align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
        set_cell_text(row_cells[1], normalize_legal_display_title(item.title), font_name=font_name)
        set_hanging_indent(row_cells[1].paragraphs[0])
        set_cell_text(row_cells[2], page_range.display(), align=WD_ALIGN_PARAGRAPH.CENTER, font_name=font_name)
        if colour_page_ranges and entry.flag_colour:
            set_cell_right_border(row_cells[2], entry.flag_colour)
    set_table_fixed_width(table, (0.55, 5.7, 1.25))


def create_authorities_index_docx(
    pdf_items: list[PDFItem],
    output_path: Path,
    *,
    font_name: str = DEFAULT_INDEX_FONT,
    index_entries: Sequence[BundleIndexEntry] | None = None,
    colour_page_ranges: bool = False,
) -> None:
    create_authorities_index_docx_from_rows(
        index_entries if index_entries is not None else get_index_entries(pdf_items),
        output_path,
        font_name=font_name,
        colour_page_ranges=colour_page_ranges,
    )


def create_authorities_index_docx_from_rows(
    index_rows: Sequence[BundleIndexRow],
    output_path: Path,
    *,
    font_name: str = DEFAULT_INDEX_FONT,
    colour_page_ranges: bool = False,
) -> None:
    doc = create_authorities_index_document(font_name=font_name)
    add_authorities_table(
        doc,
        index_rows,
        font_name=font_name,
        colour_page_ranges=colour_page_ranges,
    )
    doc.save(str(output_path))


def create_authorities_index_document(*, font_name: str = DEFAULT_INDEX_FONT) -> WordDocument:
    doc = Document()
    configure_document(doc, font_name=font_name)

    heading = doc.add_heading("Index of Authorities", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph(
        "Generated automatically from the selected PDFs. "
        "Page numbers refer to the page numbering of the combined bundle."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return doc


def add_centered_paragraph(
    doc: WordDocument,
    text: str,
    *,
    bold: bool = False,
    underline: bool = False,
    font_name: str = DEFAULT_INDEX_FONT,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.name = font_name
    run.font.size = Pt(12 if not bold else 13)


def add_label_runs(paragraph, label: str, *, font_name: str) -> None:
    if label and label[0].isdigit():
        number = "".join(ch for ch in label if ch.isdigit())
        suffix_and_role = label[len(number) :]
        suffix = suffix_and_role.split(" ", 1)[0]
        role = suffix_and_role[len(suffix) :]
        run = paragraph.add_run(number)
        run.font.name = font_name
        run.font.size = Pt(12)
        suffix_run = paragraph.add_run(suffix)
        suffix_run.font.name = font_name
        suffix_run.font.size = Pt(12)
        suffix_run.font.superscript = True
        role_run = paragraph.add_run(role)
        role_run.font.name = font_name
        role_run.font.size = Pt(12)
        return
    run = paragraph.add_run(label)
    run.font.name = font_name
    run.font.size = Pt(12)


def add_party_header_table(doc: WordDocument, matter: Matter, *, font_name: str) -> None:
    rows: list[tuple[str, str]] = []
    bringing = matter.bringing_parties
    opposing = matter.opposing_parties
    for party in bringing:
        rows.append(
            (party.name.upper(), party_label(matter.proceeding_type, PartySide.BRINGING, party.position, len(bringing)))
        )
    rows.append(("and", ""))
    for party in opposing:
        rows.append(
            (party.name.upper(), party_label(matter.proceeding_type, PartySide.OPPOSING, party.position, len(opposing)))
        )

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(PARTIES_TABLE_COLUMN_WIDTHS[0])
    table.columns[1].width = Inches(PARTIES_TABLE_COLUMN_WIDTHS[1])
    for name, label in rows:
        cells = table.add_row().cells
        name_paragraph = cells[0].paragraphs[0]
        name_run = name_paragraph.add_run(name)
        name_run.bold = name != "and"
        name_run.font.name = font_name
        name_run.font.size = Pt(12)
        label_paragraph = cells[1].paragraphs[0]
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_label_runs(label_paragraph, label, font_name=font_name)
    set_table_fixed_width(table, PARTIES_TABLE_COLUMN_WIDTHS)
    set_table_width(table, PARTIES_TABLE_WIDTH_INCHES)


def add_tramline(doc: WordDocument, *, font_name: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(PARTIES_TABLE_WIDTH_INCHES)
    set_table_width(table, PARTIES_TABLE_WIDTH_INCHES)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Pt(4)
    cell = table.cell(0, 0)
    cell.width = Inches(PARTIES_TABLE_WIDTH_INCHES)
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    set_cell_borders(cell, bottom=True)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_matter_index_header(
    doc: WordDocument,
    matter: Matter,
    document_header: DocumentHeaderInput,
    *,
    font_name: str,
) -> None:
    add_centered_paragraph(doc, matter.court.name.upper(), bold=True, font_name=font_name)
    if matter.court.header_line_2:
        add_centered_paragraph(doc, matter.court.header_line_2.upper(), font_name=font_name)

    doc.add_paragraph()
    case_paragraph = doc.add_paragraph()
    case_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    case_run = case_paragraph.add_run("CASE NO:")
    case_run.bold = True
    case_run.font.name = font_name
    case_run.font.size = Pt(12)
    if matter.case_number:
        number_run = case_paragraph.add_run(f" {matter.case_number}")
        number_run.bold = True
        number_run.font.name = font_name
        number_run.font.size = Pt(12)
    doc.add_paragraph("In the matter between:")
    doc.add_paragraph()

    add_party_header_table(doc, matter, font_name=font_name)

    doc.add_paragraph()
    add_tramline(doc, font_name=font_name)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(HEADING_TRAMLINE_SPACING_PT)
    heading.paragraph_format.space_after = Pt(HEADING_TRAMLINE_SPACING_PT)
    heading_run = heading.add_run(document_header.title.upper())
    heading_run.bold = True
    heading_run.font.name = font_name
    heading_run.font.size = Pt(13)
    add_tramline(doc, font_name=font_name)
    doc.add_paragraph()


def create_matter_authorities_index_docx(
    matter: Matter,
    document_header: DocumentHeaderInput,
    pdf_items: list[PDFItem],
    output_path: Path,
    *,
    start_page: int = 1,
    font_name: str = DEFAULT_INDEX_FONT,
    index_entries: Sequence[BundleIndexEntry] | None = None,
    colour_page_ranges: bool = False,
) -> None:
    doc = Document()
    configure_document(doc, font_name=font_name)
    add_matter_index_header(doc, matter, document_header, font_name=font_name)
    add_authorities_table(
        doc,
        index_entries if index_entries is not None else get_index_entries(pdf_items, start_page=start_page),
        font_name=font_name,
        colour_page_ranges=colour_page_ranges,
    )
    doc.save(str(output_path))


def create_matter_authorities_index_docx_from_rows(
    matter: Matter,
    document_header: DocumentHeaderInput,
    index_rows: Sequence[BundleIndexRow],
    output_path: Path,
    *,
    font_name: str = DEFAULT_INDEX_FONT,
    colour_page_ranges: bool = False,
) -> None:
    doc = Document()
    configure_document(doc, font_name=font_name)
    add_matter_index_header(doc, matter, document_header, font_name=font_name)
    add_authorities_table(
        doc,
        index_rows,
        font_name=font_name,
        colour_page_ranges=colour_page_ranges,
    )
    doc.save(str(output_path))
