#!/usr/bin/env python3

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import traceback
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog

import fitz  # PyMuPDF
from docx import Document as WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from tkinterdnd2 import DND_FILES, TkinterDnD


POSITIONS = [
    "Bottom centre",
    "Bottom right",
    "Bottom left",
    "Top centre",
    "Top right",
    "Top left",
]


@dataclass
class PDFItem:
    path: Path
    title: str


# ---------------------------------------------------------------------------
# Authority-title detection
# ---------------------------------------------------------------------------
# One PDF becomes one item in the Word index and one PDF outline entry.
# We do not extract every citation mentioned inside the authority.
# The regex is used only to infer a default title from the authority itself.

REGEX_FLAGS = re.IGNORECASE | re.VERBOSE

# Deliberately exclude forward slash. Juta source paths often look like:
#   Source: Labour Library/.../2016/May/Vox Telecommunications ...
# If slash is allowed inside party names, the regex may swallow the source path.
PARTY_CHARS = "A-Za-zÀ-ÖØ-öø-ÿ0-9&'’’.(), -"

CASE_NAME_RE = rf"""
(?:
    The[ ]+State[ ]+v[ ]+[A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,120}}
  | S[ ]+v[ ]+[A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,120}}
  | Ex[ ]+parte[ ]+[A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,170}}
  | In[ ]+re[ ]+[A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,170}}
  | [A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,170}}?[ ]+(?:v|vs[.]?)[ ]+[A-ZÀ-ÖØ-Þ][{PARTY_CHARS}]{{1,220}}?
)
"""

SA_REPORT_RE = r"""
[0-9]{4}[ ]*[(][0-9]+[)][ ]*
(?:SA|SACR|BCLR)[ ]+
[0-9]+[ ]*
[(][A-Z][A-Z0-9]{0,12}[)]
"""

ALL_SA_RE = r"""
[[][0-9]{4}[]][ ]*
[0-9]+[ ]*
All[ ]*SA[ ]+
[0-9]+[ ]*
[(][A-Z][A-Z0-9]{0,12}[)]
"""

LABOUR_REPORT_RE = r"""
(?:[0-9]{4}|[(][0-9]{4}[)])[ ]*
[0-9]+[ ]*
(?:ILJ|BLLR|BALR)[ ]+
[0-9]+[ ]*
[(][A-Z][A-Z0-9]{0,12}[)]
"""

JDR_JOL_RE = r"""
(?:[[][0-9]{4}[]]|[0-9]{4})[ ]*
(?:JDR|JOL)[ ]+
[0-9]{3,8}[ ]*
[(][A-Z][A-Z0-9]{0,12}[)]
"""

NEUTRAL_SA_RE = r"""
(?:[(][A-Za-z0-9 .,/-]+[)][ ]*)?
[[][0-9]{4}[]][.]?[ ]*
ZA[A-Z0-9]{2,14}[ ]+
[0-9]+
(?:[ ]*[(][0-9]{1,2}[ ]+[A-Za-z]+[ ]+[0-9]{4}[)])?
"""

OLD_REPORT_RE = r"""
[0-9]{4}[ ]+
(?:AD|A|NPD|TPD|CPD|OPD|WLD|GWLD|GWL|W|T|E|C|N|O|SE|Tk|Ck)[ ]+
[0-9]+
"""

ANY_CITATION_RE = rf"""
(?:
    {SA_REPORT_RE}
  | {ALL_SA_RE}
  | {LABOUR_REPORT_RE}
  | {JDR_JOL_RE}
  | {NEUTRAL_SA_RE}
  | {OLD_REPORT_RE}
)
"""

FULL_CITATION_PATTERNS = [
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{SA_REPORT_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{ALL_SA_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{LABOUR_REPORT_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{JDR_JOL_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{NEUTRAL_SA_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{CASE_NAME_RE}[ ]+{OLD_REPORT_RE})", REGEX_FLAGS),
]

LABELED_FULL_CITATION_PATTERNS = [
    re.compile(
        rf"(?:Neutral[ ]+citation|Citation)[ ]*:[ ]*(?P<citation>{CASE_NAME_RE}[ ]+{ANY_CITATION_RE})",
        REGEX_FLAGS,
    ),
]

CASE_NAME_PATTERN = re.compile(rf"(?P<name>{CASE_NAME_RE})", REGEX_FLAGS)

REPORT_ONLY_PATTERNS = [
    re.compile(rf"(?P<citation>{SA_REPORT_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{ALL_SA_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{LABOUR_REPORT_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{JDR_JOL_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{NEUTRAL_SA_RE})", REGEX_FLAGS),
    re.compile(rf"(?P<citation>{OLD_REPORT_RE})", REGEX_FLAGS),
]

BAD_CONTEXT_RE = re.compile(
    r"cases?[ ]+(?:cited|considered|referred[ ]+to)|authorities|statutes|legislation|annotations?|flynote|headnote|catchwords?|summary",
    re.IGNORECASE,
)

SECTION_BOUNDARY_RE = re.compile(
    r"(?:JUDGMENT|ORDER|REASONS|CORAM|HEARD|DELIVERED|DATE[ ]+OF[ ]+JUDGMENT|APPEARANCES|COUNSEL|ATTORNEYS)",
    re.IGNORECASE,
)

LEADING_NOISE_RE = re.compile(
    r"^(?:see also|see|cf|compare|but see|contra|in|the matter of|matter between)[ ]+",
    re.IGNORECASE,
)

TRAILING_NOISE_RE = re.compile(r"[ ,.;:]+$")

YEAR_MARKER_RE = re.compile(
    r"[[][0-9]{4}[]]|[(][0-9]{4}[)]|[0-9]{4}[ ]*[(][0-9]+[)]|[0-9]{4}[ ]+(?:AD|A|NPD|TPD|CPD|OPD|WLD|GWLD|GWL|W|T|E|C|N|O|SE|Tk|Ck)",
    re.IGNORECASE,
)

SPECIAL_WORDS = {
    "pty": "Pty",
    "ltd": "Ltd",
    "limited": "Limited",
    "inc": "Inc",
    "cc": "CC",
    "npc": "NPC",
    "no": "NO",
    "nno": "NNO",
    "mec": "MEC",
    "ccma": "CCMA",
    "numsa": "NUMSA",
    "satawu": "SATAWU",
    "nehawu": "NEHAWU",
    "sars": "SARS",
    "saps": "SAPS",
    "raf": "RAF",
    "sabc": "SABC",
    "sadtu": "SADTU",
    "cosatu": "COSATU",
    "lra": "LRA",
    "bcea": "BCEA",
    "uif": "UIF",
    "seta": "SETA",
    "soc": "SOC",
    "ndpp": "NDPP",
    "djp": "DJP",
    "aj": "AJ",
    "jp": "JP",
    "est": "Est",
    "sahrc": "SAHRC",
    "iec": "IEC",
    "icasa": "ICASA",
    "sc": "SC",
}

# Words normally lower-cased inside case names. They remain capitalised if they
# are the first word of the whole party string, e.g. "The State v ...".
LOWERCASE_WORDS = {
    "v",
    "vs",
    "and",
    "or",
    "of",
    "for",
    "in",
    "on",
    "at",
    "to",
    "by",
    "from",
    "with",
}

INITIALS_RE = re.compile(r"^(?:[A-Z][.])+$")


def smart_title_word(word: str, *, is_first_word: bool = False) -> str:
    """
    Title-case one word in a party name.

    The old version preserved any all-caps word of six letters or fewer. That
    caused false acronym preservation, for example:
        METER -> METER
        BASSON -> BASSON
        AND -> AND

    This version preserves only known acronyms/legal abbreviations and dotted
    initials. Otherwise, ordinary all-caps party names are title-cased.
    """
    if not word:
        return word

    key = word.casefold().strip(".")

    if key in SPECIAL_WORDS:
        return SPECIAL_WORDS[key]

    if INITIALS_RE.match(word):
        return word.upper()

    if key in LOWERCASE_WORDS and not is_first_word:
        return key

    if "-" in word:
        return "-".join(
            smart_title_word(part, is_first_word=is_first_word and index == 0)
            for index, part in enumerate(word.split("-"))
        )

    if "’" in word:
        return "’".join(
            smart_title_word(part, is_first_word=is_first_word and index == 0)
            for index, part in enumerate(word.split("’"))
        )

    if "'" in word:
        return "'".join(
            smart_title_word(part, is_first_word=is_first_word and index == 0)
            for index, part in enumerate(word.split("'"))
        )

    return word[:1].upper() + word[1:].lower()


def smart_title_party_text(text: str) -> str:
    parts = re.split(r"([^A-Za-zÀ-ÖØ-öø-ÿ0-9'’.-]+)", text)
    output: list[str] = []
    word_count = 0

    for index, part in enumerate(parts):
        if index % 2 == 0 and part:
            output.append(smart_title_word(part, is_first_word=(word_count == 0)))
            word_count += 1
        else:
            output.append(part)

    return "".join(output)


def titlecase_parties_before_year(citation: str) -> str:
    citation = clean_title_or_citation(citation)
    match = YEAR_MARKER_RE.search(citation)

    if not match:
        return citation

    parties = citation[: match.start()].strip()
    rest = citation[match.start() :].strip()

    if not parties:
        return citation

    return f"{smart_title_party_text(parties)} {rest}".strip()


def strip_juta_source_noise(text: str) -> str:
    """
    Remove Juta-style source/library path noise while preserving the case title.

    Example after ordinary whitespace normalisation:
        Source: Labour Library/.../2016/May/Vox Telecommunications ...

    We remove everything up to and including the final month slash, leaving:
        Vox Telecommunications ...
    """
    month_names = (
        "January|February|March|April|May|June|July|August|"
        "September|October|November|December"
    )

    text = re.sub(
        rf"Source *:.*?/[0-9]{{4}}/(?:{month_names})/",
        "",
        text,
        flags=re.IGNORECASE,
    )

    text = re.sub(
        rf"Source *:.*?Chronological +listing +[0-9]{{4}} +(?:{month_names}) +",
        "",
        text,
        flags=re.IGNORECASE,
    )

    return text


def normalise_pdf_text(text: str) -> str:
    text = text.replace(chr(173), "")
    text = re.sub(
        r"([A-Za-zÀ-ÖØ-öø-ÿ])- +([A-Za-zÀ-ÖØ-öø-ÿ])",
        lambda m: m.group(1) + m.group(2),
        text,
    )
    text = " ".join(text.split())
    text = strip_juta_source_noise(text)
    return text


def clean_title_or_citation(text: str) -> str:
    text = text.replace(chr(173), "")
    text = " ".join(text.split()).strip()
    text = re.sub(r" +([,.;:])", lambda m: m.group(1), text)
    text = re.sub(r"[(] +", "(", text)
    text = re.sub(r" +[)]", ")", text)
    text = LEADING_NOISE_RE.sub("", text)
    text = TRAILING_NOISE_RE.sub("", text)
    return text.strip()


def clean_filename_title(path: Path) -> str:
    title = path.stem
    title = re.sub(r"[_]+", " ", title)
    title = re.sub(r" +- +", " - ", title)
    title = " ".join(title.split()).strip()
    return title or path.name


def top_text_from_first_page(doc: fitz.Document, top_fraction: float = 0.60) -> str:
    if doc.page_count == 0:
        return ""

    page = doc[0]
    cutoff_y = page.rect.y0 + page.rect.height * top_fraction
    blocks = page.get_text("blocks", sort=True)
    parts: list[str] = []

    for block in blocks:
        x0, y0, x1, y1, text = block[:5]
        if y0 <= cutoff_y:
            parts.append(str(text))

    return normalise_pdf_text(chr(10).join(parts))


def first_pages_text(doc: fitz.Document, pages: int = 2) -> str:
    parts: list[str] = []
    for i in range(min(pages, doc.page_count)):
        parts.append(doc[i].get_text("text", sort=True))
    return normalise_pdf_text(chr(10).join(parts))


def text_before_first_boundary(text: str) -> str:
    match = SECTION_BOUNDARY_RE.search(text)
    if not match:
        return text
    return text[: match.start()]


def candidate_penalty_for_context(text: str, start: int) -> int:
    context = text[max(0, start - 160) : start]
    return -140 if BAD_CONTEXT_RE.search(context) else 0


def find_first_report_after(text: str, start: int, window: int = 600) -> str | None:
    snippet = text[start : start + window]
    best: tuple[int, str] | None = None

    for pattern in REPORT_ONLY_PATTERNS:
        for match in pattern.finditer(snippet):
            citation = clean_title_or_citation(match.group("citation"))
            if not citation:
                continue
            if best is None or match.start() < best[0]:
                best = (match.start(), citation)

    if best is None:
        return None

    return best[1]


def find_best_full_citation_candidate(text: str, base_score: int) -> tuple[int, str] | None:
    best: tuple[int, str] | None = None

    for pattern in LABELED_FULL_CITATION_PATTERNS:
        for match in pattern.finditer(text):
            candidate = titlecase_parties_before_year(match.group("citation"))
            if not candidate:
                continue
            score = base_score + 160 + candidate_penalty_for_context(text, match.start())
            score -= min(match.start() // 200, 25)
            if best is None or score > best[0]:
                best = (score, candidate)

    for pattern in FULL_CITATION_PATTERNS:
        for match in pattern.finditer(text):
            candidate = titlecase_parties_before_year(match.group("citation"))
            if not candidate:
                continue
            score = base_score + 90 + candidate_penalty_for_context(text, match.start())
            score -= min(match.start() // 120, 45)
            if best is None or score > best[0]:
                best = (score, candidate)

    return best


def find_case_name_plus_nearby_report(text: str) -> tuple[int, str] | None:
    for match in CASE_NAME_PATTERN.finditer(text):
        if candidate_penalty_for_context(text, match.start()) < 0:
            continue

        name = clean_title_or_citation(match.group("name"))
        if not name or len(name) > 260:
            continue

        report = find_first_report_after(text, match.end())
        if not report:
            continue

        return match.start(), titlecase_parties_before_year(f"{name} {report}")

    return None


def detect_authority_index_item(path: Path) -> str:
    """
    Best-effort detection of the PDF's own full case citation.

    If no full citation is found, fall back to the filename. We do not return
    every cited case inside the document, and we do not return a bare case name
    without a citation as the default.
    """
    fallback = clean_filename_title(path)

    try:
        doc = fitz.open(path)
    except Exception:
        return fallback

    try:
        if doc.page_count == 0:
            return fallback

        top_first_page = top_text_from_first_page(doc)
        first_two_pages = first_pages_text(doc, pages=2)
        early_text = text_before_first_boundary(first_two_pages)
        title_text = top_first_page or early_text[:3500]

        candidates: list[tuple[int, str]] = []

        for text, base_score in [
            (title_text, 90),
            (early_text, 55),
            (first_two_pages[:5500], 15),
        ]:
            full = find_best_full_citation_candidate(text, base_score)
            if full is not None:
                candidates.append(full)

            nearby = find_case_name_plus_nearby_report(text)
            if nearby is not None:
                start, citation = nearby
                score = base_score + 60 + candidate_penalty_for_context(text, start)
                score -= min(start // 120, 45)
                candidates.append((score, citation))

        if candidates:
            candidates.sort(key=lambda item: item[0], reverse=True)
            return candidates[0][1]

        return fallback
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# Word index output
# ---------------------------------------------------------------------------

def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, bold: bool = False, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def create_authorities_index_docx(pdf_items: list[PDFItem], output_path: Path) -> None:
    doc = WordDocument()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    heading = doc.add_heading("Index of Authorities", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph(
        "Generated automatically from the selected PDFs. Page numbers refer to the page numbering of the combined bundle."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(5.7)
    table.columns[2].width = Inches(1.25)

    header_cells = table.rows[0].cells
    set_cell_text(header_cells[0], "No", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_cells[1], "Item", bold=True)
    set_cell_text(header_cells[2], "Page no", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    global_page_no = 1

    for number, item in enumerate(pdf_items, start=1):
        row_cells = table.add_row().cells
        set_cell_text(row_cells[0], str(number), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row_cells[1], item.title)
        set_cell_text(row_cells[2], str(global_page_no), align=WD_ALIGN_PARAGRAPH.CENTER)

        try:
            source = fitz.open(item.path)
            global_page_no += source.page_count
            source.close()
        except Exception:
            global_page_no += 1

    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF combining / page numbering / PDF outline
# ---------------------------------------------------------------------------

def get_number_box(page_rect: fitz.Rect, position: str, box_width: float, box_height: float, margin: float) -> fitz.Rect:
    position_lower = position.lower()

    if "left" in position_lower:
        x0 = page_rect.x0 + margin
    elif "right" in position_lower:
        x0 = page_rect.x1 - margin - box_width
    else:
        x0 = page_rect.x0 + (page_rect.width - box_width) / 2

    if "top" in position_lower:
        y0 = page_rect.y0 + margin
    else:
        y0 = page_rect.y1 - margin - box_height

    return fitz.Rect(x0, y0, x0 + box_width, y0 + box_height)


def draw_page_number(page: fitz.Page, number: int, position: str, font_size: int = 15, margin: int = 28) -> None:
    text = str(number)
    page_rect = page.rect

    text_width = fitz.get_text_length(text, fontname="hebo", fontsize=font_size)
    box_width = max(34, text_width + 18)
    box_height = font_size + 12

    box = get_number_box(page_rect, position, box_width, box_height, margin)

    page.draw_rect(
        box,
        color=(0, 0, 0),
        fill=(1, 1, 1),
        width=0.6,
        fill_opacity=0.85,
        stroke_opacity=0.65,
    )

    page.insert_textbox(
        box,
        text,
        fontsize=font_size,
        fontname="hebo",
        color=(0, 0, 0),
        align=fitz.TEXT_ALIGN_CENTER,
    )


def combine_number_and_add_toc(pdf_items: list[PDFItem], output_path: Path, position: str, font_size: int, margin: int) -> None:
    if not pdf_items:
        raise ValueError("No PDFs selected.")

    output_doc = fitz.open()
    toc: list[list[int | str]] = []

    try:
        current_start_page = 1

        for item in pdf_items:
            source = fitz.open(item.path)
            try:
                if source.page_count == 0:
                    continue

                toc_title = item.title.strip() or clean_filename_title(item.path)
                toc.append([1, toc_title, current_start_page])

                output_doc.insert_pdf(source)
                current_start_page += source.page_count
            finally:
                source.close()

        for index, page in enumerate(output_doc, start=1):
            draw_page_number(page, index, position, font_size, margin)

        if toc:
            output_doc.set_toc(toc)

        output_doc.save(output_path, garbage=4, deflate=True)
    finally:
        output_doc.close()


# ---------------------------------------------------------------------------
# Tkinter app
# ---------------------------------------------------------------------------

class PDFCombinerNumbererTOCIndexApp(TkinterDnD.Tk):
    def __init__(self) -> None:
        super().__init__()

        self.title("PDF Combiner + Page Numberer + Authorities Index")
        self.geometry("830x540")
        self.minsize(760, 480)

        self.pdf_items: list[PDFItem] = []

        self.position_var = tk.StringVar(value="Bottom centre")
        self.font_size_var = tk.IntVar(value=15)
        self.margin_var = tk.IntVar(value=28)

        self._build_ui()

    def _build_ui(self) -> None:
        outer = ttk.Frame(self, padding=16)
        outer.pack(fill="both", expand=True)

        title = ttk.Label(
            outer,
            text="Drag PDFs here, edit titles, then create a numbered bundle or authorities index",
            font=("TkDefaultFont", 15, "bold"),
        )
        title.pack(anchor="w", pady=(0, 12))

        controls = ttk.Frame(outer)
        controls.pack(fill="x", pady=(0, 12))

        ttk.Label(controls, text="Number position:").grid(row=0, column=0, sticky="w")

        position_box = ttk.Combobox(
            controls,
            textvariable=self.position_var,
            values=POSITIONS,
            state="readonly",
            width=20,
        )
        position_box.grid(row=0, column=1, sticky="w", padx=(8, 24))

        ttk.Label(controls, text="Font size:").grid(row=0, column=2, sticky="w")

        font_spin = ttk.Spinbox(
            controls,
            from_=10,
            to=28,
            textvariable=self.font_size_var,
            width=5,
        )
        font_spin.grid(row=0, column=3, sticky="w", padx=(8, 24))

        ttk.Label(controls, text="Margin:").grid(row=0, column=4, sticky="w")

        margin_spin = ttk.Spinbox(
            controls,
            from_=10,
            to=100,
            textvariable=self.margin_var,
            width=5,
        )
        margin_spin.grid(row=0, column=5, sticky="w", padx=(8, 0))

        main = ttk.Frame(outer)
        main.pack(fill="both", expand=True)

        left = ttk.Frame(main)
        left.pack(side="left", fill="both", expand=True)

        self.tree = ttk.Treeview(
            left,
            columns=("order", "title"),
            show="headings",
            selectmode="browse",
        )

        self.tree.heading("order", text="#")
        self.tree.heading("title", text="PDF ToC entry / index item")

        self.tree.column("order", width=45, anchor="center", stretch=False)
        self.tree.column("title", width=650, anchor="w", stretch=True)

        self.tree.pack(fill="both", expand=True)

        self.tree.drop_target_register(DND_FILES)
        self.tree.dnd_bind("<<Drop>>", self.on_drop)

        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.tree.bind("<Return>", lambda event: self.edit_selected_title())

        hint = ttk.Label(
            left,
            text="Double-click a title to edit it. This title is used for both the PDF bookmark and the Word index row.",
        )
        hint.pack(anchor="w", pady=(6, 0))

        buttons = ttk.Frame(main)
        buttons.pack(side="right", fill="y", padx=(12, 0))

        ttk.Button(buttons, text="Add PDFs", command=self.add_pdfs_dialog).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Edit title", command=self.edit_selected_title).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Auto-detect title", command=self.auto_detect_selected_title).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Move up", command=self.move_up).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Move down", command=self.move_down).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Remove", command=self.remove_selected).pack(fill="x", pady=(0, 6))
        ttk.Button(buttons, text="Clear", command=self.clear_list).pack(fill="x", pady=(0, 18))

        ttk.Button(
            buttons,
            text="Create combined numbered PDF",
            command=self.create_combined_pdf,
        ).pack(fill="x", pady=(0, 6))

        ttk.Button(
            buttons,
            text="Create authorities index (.docx)",
            command=self.create_authorities_index,
        ).pack(fill="x")

        self.status = ttk.Label(outer, text="No PDFs added yet.", anchor="w")
        self.status.pack(fill="x", pady=(12, 0))

    def parse_drop_files(self, data: str) -> list[Path]:
        raw_paths = self.tk.splitlist(data)
        return [Path(p) for p in raw_paths if Path(p).is_file() and Path(p).suffix.lower() == ".pdf"]

    def add_paths(self, paths: list[Path]) -> None:
        existing = {item.path.resolve() for item in self.pdf_items}
        added = 0

        for path in paths:
            resolved = path.resolve()

            if resolved in existing:
                continue

            detected_title = detect_authority_index_item(path)

            self.pdf_items.append(PDFItem(path=path, title=detected_title))
            existing.add(resolved)
            added += 1

        self.refresh_tree()

        if added:
            self.status.config(text=f"Added {added} PDF(s). Total: {len(self.pdf_items)}.")
        else:
            self.status.config(text="No new PDFs added.")

    def refresh_tree(self) -> None:
        self.tree.delete(*self.tree.get_children())

        for index, item in enumerate(self.pdf_items, start=1):
            self.tree.insert("", tk.END, iid=str(index - 1), values=(index, item.title))

    def selected_index(self) -> int | None:
        selection = self.tree.selection()
        if not selection:
            return None
        return int(selection[0])

    def on_drop(self, event) -> None:
        pdfs = self.parse_drop_files(event.data)

        if not pdfs:
            messagebox.showwarning("No PDFs", "Please drop one or more PDF files.")
            return

        self.add_paths(pdfs)

    def add_pdfs_dialog(self) -> None:
        selected = filedialog.askopenfilenames(title="Choose PDFs", filetypes=[("PDF files", "*.pdf")])
        if selected:
            self.add_paths([Path(p) for p in selected])

    def on_tree_double_click(self, event) -> None:
        region = self.tree.identify("region", event.x, event.y)
        column = self.tree.identify_column(event.x)

        if region == "cell" and column == "#2":
            self.edit_selected_title()

    def edit_selected_title(self) -> None:
        index = self.selected_index()

        if index is None:
            messagebox.showwarning("Nothing selected", "Please select a PDF first.")
            return

        item = self.pdf_items[index]

        new_title = simpledialog.askstring(
            title="Edit title",
            prompt="PDF ToC entry / index item:",
            initialvalue=item.title,
            parent=self,
        )

        if new_title is None:
            return

        new_title = new_title.strip()

        if not new_title:
            messagebox.showwarning("Blank title", "The title cannot be blank.")
            return

        item.title = new_title
        self.refresh_tree()
        self.tree.selection_set(str(index))
        self.status.config(text=f"Renamed title to: {new_title}")

    def auto_detect_selected_title(self) -> None:
        index = self.selected_index()

        if index is None:
            messagebox.showwarning("Nothing selected", "Please select a PDF first.")
            return

        item = self.pdf_items[index]
        item.title = detect_authority_index_item(item.path)
        self.refresh_tree()
        self.tree.selection_set(str(index))
        self.status.config(text=f"Auto-detected title: {item.title}")

    def move_up(self) -> None:
        index = self.selected_index()
        if index is None or index == 0:
            return

        self.pdf_items[index - 1], self.pdf_items[index] = self.pdf_items[index], self.pdf_items[index - 1]
        self.refresh_tree()
        self.tree.selection_set(str(index - 1))

    def move_down(self) -> None:
        index = self.selected_index()
        if index is None or index >= len(self.pdf_items) - 1:
            return

        self.pdf_items[index + 1], self.pdf_items[index] = self.pdf_items[index], self.pdf_items[index + 1]
        self.refresh_tree()
        self.tree.selection_set(str(index + 1))

    def remove_selected(self) -> None:
        index = self.selected_index()
        if index is None:
            return

        removed = self.pdf_items.pop(index)
        self.refresh_tree()
        self.status.config(text=f"Removed {removed.path.name}.")

    def clear_list(self) -> None:
        self.pdf_items.clear()
        self.refresh_tree()
        self.status.config(text="List cleared.")

    def create_combined_pdf(self) -> None:
        if not self.pdf_items:
            messagebox.showwarning("No PDFs", "Please add one or more PDFs first.")
            return

        output_path_str = filedialog.asksaveasfilename(
            title="Save combined numbered PDF as",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile="combined_numbered.pdf",
        )

        if not output_path_str:
            return

        output_path = Path(output_path_str)

        try:
            combine_number_and_add_toc(
                pdf_items=self.pdf_items,
                output_path=output_path,
                position=self.position_var.get(),
                font_size=int(self.font_size_var.get()),
                margin=int(self.margin_var.get()),
            )
        except Exception as exc:
            traceback.print_exc()
            messagebox.showerror("Failed", f"Could not create PDF: {exc}")
            self.status.config(text="Failed to create combined PDF.")
            return

        self.status.config(text=f"Created: {output_path}")
        messagebox.showinfo("Complete", f"Created: {output_path}")

    def create_authorities_index(self) -> None:
        if not self.pdf_items:
            messagebox.showwarning("No PDFs", "Please add one or more PDFs first.")
            return

        output_path_str = filedialog.asksaveasfilename(
            title="Save authorities index as",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile="authorities_index.docx",
        )

        if not output_path_str:
            return

        output_path = Path(output_path_str)
        self.status.config(text="Creating authorities index...")
        self.update_idletasks()

        try:
            create_authorities_index_docx(self.pdf_items, output_path)
        except Exception as exc:
            traceback.print_exc()
            messagebox.showerror("Failed", f"Could not create authorities index: {exc}")
            self.status.config(text="Failed to create authorities index.")
            return

        self.status.config(text=f"Created authorities index: {output_path}")
        messagebox.showinfo("Complete", f"Created authorities index: {output_path}")


def main() -> None:
    app = PDFCombinerNumbererTOCIndexApp()
    app.mainloop()


if __name__ == "__main__":
    main()
